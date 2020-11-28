import time
import docx
import re
#python-docx将整个文章看做是一个Document对象 官方文档 - Document，其基本结构如下：
#每个Document包含许多个代表“段落”的Paragraph对象，存放在document.paragraphs中。
#每个Paragraph都有许多个代表"行内元素"的Run对象，存放在paragraph.runs中。
#在python-docx中，run是最基本的单位，每个run对象内的文本样式都是一致的
# 在从docx文件生成文档对象时，python-docx会根据样式的变化来将文本切分为一个个的Run对象。

def diff(listA,listB):
    #求交集的两种方式
    retA = [i for i in listA if i in listB]
    retB = list(set(listA).intersection(set(listB)))
    
    print("retA is: ",retA)
    print("retB is: ",retB)
    
    #求并集
    retC = list(set(listA).union(set(listB)))
    print("retC1 is: ",retC)
    
    #求差集，在B中但不在A中
    retD = list(set(listB).difference(set(listA)))
    print("差集，在B中但不在A中 retD is: ",retD)
    
    retE = [i for i in listB if i not in listA]
    print("差集，在B中但不在A中 retE is: ",retE)

def replace_texts(inPathName, old_texts, new_texts):
    fdocx = docx.Document(inPathName)
    '''此函数用于批量替换合同中需要替换的信息
    doc:文件
    old_info和new_info：原文字和需要替换的新文字
    '''

    #读取段落中的所有run，找到需替换的信息进行替换
    signal = False
    for para in fdocx.paragraphs: #
        #print("*"*20,"\npara.text ",para.text,"\n","*"*20)
        for i in old_texts:
            if i in para:
                signal = signal or True
        if signal:
            for run in para.runs:
                print("[run]: ",run.text)
                for i in range(0,len(old_texts)):
                    # 使用 runs 替换内容但不改变样式
                    # 注意！runs 会根据样式分隔内容，确保被替换内容的样式一致
                    if run.text: #避免图片丢失
                        run.text = run.text.replace(old_texts[i], new_texts[i]) #替换信息

    #读取表格中的所有单元格，找到需替换的信息进行替换
    signal = False
    for table in fdocx.tables:
        for row in table.rows:
            for cell in row.cells:
                #print("-*"*20,"\n",cell.text,"\n","*-"*20,"\n")
                for para in cell.paragraphs:
                    for i in old_texts:
                        if i in para:
                            signal = signal or True
                    if signal:
                        for run in para.runs:
                            for i in range(0,len(old_texts)):
                                if run.text: #避免图片丢失
                                    run.text = run.text.replace(old_texts[i], new_texts[i]) #替换信息
    fdocx.save(inPathName.split(".docx")[0]+"_"+str(time.time())+"_xiugai.docx")
    

def gettexts(inPathNames):
    fdocxs = []#document对象列表
    fdocxtexts = []#document对象内短句列表
    for filename in inPathNames:
        fdocxs.append(docx.Document(filename))
    for fdocx in fdocxs:
        texts = []
        #读取段落中的所有text
        for para in fdocx.paragraphs:
            if para.text:
                texts += re.split(r'\s*([;,\.\!\?；，。？！]+)\s*', para.text)
        #读取表格中的所有text
        for table in fdocx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        texts += re.split(r'\s*([;,\.\!\?；，。？！]+)\s*', para.text)
        for i in ["，", "；", "。", "！", "？", "", ",", ";", ".", "!", "?"]:
            try:
                texts.remove(i)
            except Exception as e:
                print(e)
                pass
        fdocxtexts.append(texts)
    for texts in fdocxtexts:
        print(texts,"\n")

def find_same_twof(inPathNames):
    fdocxs = []#document对象列表
    fdocxtexts = []#document对象内短句列表
    for filename in inPathNames:
        fdocxs.append(docx.Document(filename))
    for fdocx in fdocxs:
        texts = []
        #读取段落中的所有text
        for para in fdocx.paragraphs:
            if para.text:
                texts += re.split(r'\s*([;,\.\!\?；，。？！]+)\s*', para.text)
        #读取表格中的所有text
        for table in fdocx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        texts += re.split(r'\s*([;,\.\!\?；，。？！]+)\s*', para.text)
        for i in ["，", "；", "。", "！", "？", "", ",", ";", ".", "!", "?"]:
            try:
                texts.remove(i)
            except Exception as e:
                print(e)
                pass
        fdocxtexts.append(texts)
    print(fdocxtexts[0],"\n")
    print(fdocxtexts[1],"\n")
    if fdocxtexts[2]:
        retB = [i for i in fdocxtexts[0] if i in fdocxtexts[1] and i not in fdocxtexts[2]]
    else:
        retB = [i for i in fdocxtexts[0] if i in fdocxtexts[1]]
    #retB = list(set(retB))
    for texts in retB:
        print(texts)

#replace_texts("test.docx",["文件"],["【 FILE 】"])

#gettexts(["test.docx",'test_2.docx'])
pathtemp= "/Users/liudyixia/OneDrive/瑞数信息/20181022四川电信企信部门/四川电信投标20191230-20200113"
A = pathtemp + '四川电信技术文件-瑞数.docx'
B = pathtemp + '2019年四川电信投标文件1.4.docx'
find_same_twof([A, B])
'''
fpath = "test.docx"
file=docx.Document(fpath)
print('段落:'+str(len(file.paragraphs)))
for i in range(len(file.paragraphs)):
    print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)
'''